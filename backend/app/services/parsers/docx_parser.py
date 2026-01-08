"""DOCX parser using python-docx."""
import logging
from pathlib import Path

from docx import Document

from app.services.parsers.base_parser import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX files."""
    
    SUPPORTED_EXTENSIONS = {'.docx', '.doc'}
    
    def can_parse(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> ParseResult:
        """Parse DOCX and extract text."""
        self._log_parse_start(file_path)
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n".join(paragraphs)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text += "\n" + row_text
            
            if not full_text.strip():
                raise ValueError("No text could be extracted from this DOCX file.")
            
            # Estimate page count (rough: ~500 words per page)
            word_count = len(full_text.split())
            page_count = max(1, word_count // 500)
            
            result = ParseResult(
                text=full_text,
                page_count=page_count,
                word_count=word_count,
                metadata={"parser": "python-docx", "file_path": file_path}
            )
            
            self._log_parse_complete(file_path, result)
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
